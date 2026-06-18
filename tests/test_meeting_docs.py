import importlib.util
import unittest

from whisper_live.meeting.docs import MeetingDocConverter


@unittest.skipUnless(importlib.util.find_spec("docx"), "python-docx is not installed")
class TestMeetingDocConverter(unittest.TestCase):
    def test_md_to_docx_bytes_converts_basic_markdown(self):
        content = "\n".join([
            "# 会议纪要",
            "",
            "## 会议议题",
            "- **领域决赛**安排",
            "1. 后续处理",
            "| 项目 | 内容 |",
            "| --- | --- |",
            "| 会议名称 | 周会 |",
        ])

        data = MeetingDocConverter.md_to_docx_bytes(content)

        self.assertTrue(data.startswith(b"PK"))
        self.assertGreater(len(data), 1000)

    def test_md_to_docx_applies_official_fonts(self):
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
        import io

        content = "\n".join([
            "# 会议纪要",
            "正文内容",
            "| 项目 | 内容 |",
            "| --- | --- |",
            "| 会议名称 | 周会 |",
        ])
        data = MeetingDocConverter.md_to_docx_bytes(content)
        document = Document(io.BytesIO(data))

        normal = document.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")
        self.assertEqual(normal.element.rPr.rFonts.get(qn("w:eastAsia")), "SimSun")
        self.assertEqual(normal.font.size, Pt(12))

        heading = document.styles["Heading 1"]
        self.assertEqual(heading.font.name, "Times New Roman")
        self.assertEqual(heading.element.rPr.rFonts.get(qn("w:eastAsia")), "SimSun")
        self.assertTrue(heading.font.bold)
        self.assertEqual(heading.font.size, Pt(16))

        table_run = document.tables[0].cell(0, 0).paragraphs[0].runs[0]
        self.assertEqual(table_run.font.name, "Times New Roman")
        self.assertEqual(table_run._element.rPr.rFonts.get(qn("w:eastAsia")), "SimSun")
        self.assertTrue(table_run.bold)
        self.assertEqual(table_run.font.size, Pt(12))

    def test_docx_to_md_text_round_trips_basic_document(self):
        from docx import Document
        import tempfile
        import os

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            document = Document()
            document.add_heading("会议纪要", level=1)
            document.add_paragraph("正文内容")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "内容"
            document.save(path)

            text = MeetingDocConverter.docx_to_md_text(path)

        self.assertIn("# 会议纪要", text)
        self.assertIn("正文内容", text)
        self.assertIn("| 项目 | 内容 |", text)

    def test_docx_to_md_text_promotes_plain_meeting_template_headings(self):
        from docx import Document
        import tempfile
        import os

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "plain-template.docx")
            document = Document()
            document.add_paragraph("会议纪要")
            document.add_paragraph("会议基本信息")
            document.add_paragraph("模板示例内容")
            document.add_paragraph("一、会议议题")
            document.add_paragraph("- 示例议题")
            document.add_paragraph("讨论事项综述")
            document.save(path)

            text = MeetingDocConverter.docx_to_md_text(path, promote_plain_headings=True)

        self.assertIn("# 会议纪要", text)
        self.assertIn("## 会议基本信息", text)
        self.assertIn("## 会议议题", text)
        self.assertIn("## 讨论事项综述", text)
        self.assertIn("模板示例内容", text)

    def test_docx_to_md_text_does_not_promote_plain_body_text(self):
        from docx import Document
        import tempfile
        import os

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "body.docx")
            document = Document()
            document.add_paragraph("这是普通正文")
            document.add_paragraph("项目按计划推进")
            document.save(path)

            text = MeetingDocConverter.docx_to_md_text(path, promote_plain_headings=True)

        self.assertNotIn("## 这是普通正文", text)
        self.assertNotIn("## 项目按计划推进", text)
        self.assertIn("这是普通正文", text)
        self.assertIn("项目按计划推进", text)
