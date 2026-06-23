import io
import os
import re


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class MeetingDocConverter:
    """Small Markdown/DOCX converter for meeting summaries."""

    DEFAULT_EAST_ASIA_FONT = "SimSun"
    DEFAULT_LATIN_FONT = "Times New Roman"
    DEFAULT_BODY_SIZE_PT = 12
    HEADING_SIZES_PT = {1: 16, 2: 14, 3: 12, 4: 12}
    PLAIN_DOCUMENT_TITLES = {"会议纪要", "会议记录", "会议总结"}
    PLAIN_FIELD_HEADINGS = {
        "会议基本信息": ("会议基本信息", "会议信息", "基本信息"),
        "参会人员": ("参会人员", "出席人员", "与会人员"),
        "会议议题": ("会议议题", "会议议程", "议题"),
        "讨论事项综述": ("讨论事项综述", "讨论事项", "事项综述", "会议内容", "讨论内容"),
        "会议决议": ("会议决议", "决议事项", "会议结论", "结论"),
        "待办事项": ("待办事项", "行动项", "后续事项", "下一步计划"),
        "问题与风险": ("问题与风险", "风险问题", "遗留问题"),
    }


    @staticmethod
    def _document_class():
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "Missing dependency: python-docx. Install requirements/docx.txt to enable DOCX export."
            ) from exc
        return Document

    @staticmethod
    def _docx_units():
        try:
            from docx.enum.text import WD_LINE_SPACING
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except ImportError as exc:
            raise RuntimeError(
                "Missing dependency: python-docx. Install requirements/docx.txt to enable DOCX export."
            ) from exc
        return Cm, Pt, WD_LINE_SPACING, qn

    @staticmethod
    def _normalize_plain_heading(text):
        value = re.sub(r"^[\s#]*", "", str(text or "").strip())
        value = re.sub(r"^(第?[一二三四五六七八九十百千0-9]+[章节部分项]?[、.．)）:：\-\s]+)", "", value)
        value = value.strip(" ：:;；。.-—\t")
        return re.sub(r"\s+", "", value)

    @classmethod
    def _plain_heading_markdown(cls, text):
        normalized = cls._normalize_plain_heading(text)
        if normalized in cls.PLAIN_DOCUMENT_TITLES:
            return f"# {normalized}"
        for canonical, aliases in cls.PLAIN_FIELD_HEADINGS.items():
            if normalized in aliases:
                return f"## {canonical}"
        return None

    @staticmethod
    def _docx_heading_level(paragraph):
        style = paragraph.style
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            name = str(getattr(style, "name", "") or "").strip()
            match = re.match(r"^(?:Heading|标题)\s*([1-6])$", name, re.IGNORECASE)
            if match:
                return int(match.group(1))
            ppr = getattr(getattr(style, "element", None), "pPr", None)
            outline = getattr(ppr, "outlineLvl", None) if ppr is not None else None
            if outline is not None and getattr(outline, "val", None) is not None:
                try:
                    return max(1, min(int(outline.val) + 1, 6))
                except (TypeError, ValueError):
                    pass
            style = getattr(style, "base_style", None)
        paragraph_style_name = str(getattr(paragraph.style, "name", "") or "").strip().lower()
        if paragraph_style_name not in {"normal", "正文", "body text"}:
            ppr = getattr(paragraph._p, "pPr", None)
            outline = getattr(ppr, "outlineLvl", None) if ppr is not None else None
            if outline is not None and getattr(outline, "val", None) is not None:
                try:
                    return max(1, min(int(outline.val) + 1, 6))
                except (TypeError, ValueError):
                    pass
        return None

    @staticmethod
    def _docx_run_signature(paragraph):
        try:
            from docx.oxml.ns import qn
        except ImportError:
            return set(), False
        fonts = set()
        bold = False
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            if run.font.name:
                fonts.add(str(run.font.name).lower())
            rpr = run._element.rPr
            rfonts = rpr.rFonts if rpr is not None else None
            if rfonts is not None:
                east_asia = rfonts.get(qn("w:eastAsia"))
                if east_asia:
                    fonts.add(str(east_asia).lower())
            bold = bold or bool(run.bold)
        return fonts, bold

    @classmethod
    def _is_plain_numbered_subheading(cls, paragraph, next_paragraph=None):
        text = paragraph.text.strip()
        match = re.match(r"^(\d{1,2})\s*[.、．)）]\s*(\S.+?)\s*$", text)
        if not match:
            return False
        title = match.group(2).strip()
        if not 2 <= len(title) <= 100 or re.search(r"[。！？!?；;]$", title):
            return False
        title_markers = (
            "安排", "汇报", "讨论", "情况", "方案", "事项", "计划", "进展",
            "处置", "办法", "更新", "总结", "问题", "审议", "验收", "建设",
        )
        keyword_match = title.endswith(title_markers)
        fonts, bold = cls._docx_run_signature(paragraph)
        next_fonts, _next_bold = cls._docx_run_signature(next_paragraph) if next_paragraph is not None else (set(), False)
        format_differs = bool(fonts and next_fonts and fonts != next_fonts)
        return bold or format_differs or keyword_match

    @staticmethod
    def _split_table_row(line):
        body = line.strip().strip("|")
        return [cell.strip() for cell in body.split("|")]

    @staticmethod
    def _is_table_separator(line):
        cells = MeetingDocConverter._split_table_row(line)
        return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)

    @staticmethod
    def _apply_run_font(run, size_pt=None, bold=None):
        _Cm, Pt, _WD_LINE_SPACING, qn = MeetingDocConverter._docx_units()
        run.font.name = MeetingDocConverter.DEFAULT_LATIN_FONT
        run.font.size = Pt(size_pt or MeetingDocConverter.DEFAULT_BODY_SIZE_PT)
        if bold is not None:
            run.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), MeetingDocConverter.DEFAULT_EAST_ASIA_FONT)
        rfonts.set(qn("w:ascii"), MeetingDocConverter.DEFAULT_LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), MeetingDocConverter.DEFAULT_LATIN_FONT)

    @staticmethod
    def _apply_paragraph_font(paragraph, size_pt=None, bold=None):
        for run in paragraph.runs:
            MeetingDocConverter._apply_run_font(run, size_pt=size_pt, bold=bold)

    @staticmethod
    def _configure_style_font(style, size_pt=None, bold=None):
        _Cm, Pt, _WD_LINE_SPACING, qn = MeetingDocConverter._docx_units()
        style.font.name = MeetingDocConverter.DEFAULT_LATIN_FONT
        style.font.size = Pt(size_pt or MeetingDocConverter.DEFAULT_BODY_SIZE_PT)
        if bold is not None:
            style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), MeetingDocConverter.DEFAULT_EAST_ASIA_FONT)
        rfonts.set(qn("w:ascii"), MeetingDocConverter.DEFAULT_LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), MeetingDocConverter.DEFAULT_LATIN_FONT)

    @staticmethod
    def _configure_document_styles(document):
        Cm, Pt, WD_LINE_SPACING, _qn = MeetingDocConverter._docx_units()
        for section in document.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        styles = document.styles
        MeetingDocConverter._configure_style_font(styles["Normal"], MeetingDocConverter.DEFAULT_BODY_SIZE_PT)
        styles["Normal"].paragraph_format.space_after = Pt(4)
        styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        styles["Normal"].paragraph_format.line_spacing = 1.15

        for level, size in MeetingDocConverter.HEADING_SIZES_PT.items():
            style_name = f"Heading {level}"
            if style_name in styles:
                MeetingDocConverter._configure_style_font(styles[style_name], size_pt=size, bold=True)
                styles[style_name].paragraph_format.space_before = Pt(8)
                styles[style_name].paragraph_format.space_after = Pt(6)

        for style_name in ("List Bullet", "List Number"):
            if style_name in styles:
                MeetingDocConverter._configure_style_font(styles[style_name], MeetingDocConverter.DEFAULT_BODY_SIZE_PT)
                styles[style_name].paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_inline_markdown(paragraph, text, size_pt=None, bold=None):
        text = str(text or "")
        pos = 0
        for match in re.finditer(r"\*\*(.+?)\*\*", text):
            if match.start() > pos:
                run = paragraph.add_run(text[pos:match.start()])
                MeetingDocConverter._apply_run_font(run, size_pt=size_pt, bold=bold)
            run = paragraph.add_run(match.group(1))
            MeetingDocConverter._apply_run_font(run, size_pt=size_pt, bold=True)
            pos = match.end()
        if pos < len(text):
            run = paragraph.add_run(text[pos:])
            MeetingDocConverter._apply_run_font(run, size_pt=size_pt, bold=bold)

    @staticmethod
    def _add_paragraph(document, text, style=None, size_pt=None, bold=None):
        paragraph = document.add_paragraph(style=style)
        MeetingDocConverter._add_inline_markdown(paragraph, text, size_pt=size_pt, bold=bold)
        return paragraph

    @staticmethod
    def _add_table(document, rows):
        if not rows:
            return
        width = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for col_index in range(width):
                cell = table.cell(row_index, col_index)
                paragraph = cell.paragraphs[0]
                paragraph._p.clear_content()
                run = paragraph.add_run(row[col_index] if col_index < len(row) else "")
                MeetingDocConverter._apply_run_font(
                    run,
                    size_pt=MeetingDocConverter.DEFAULT_BODY_SIZE_PT,
                    bold=(row_index == 0),
                )

    @staticmethod
    def md_to_docx_bytes(md_content: str) -> bytes:
        Document = MeetingDocConverter._document_class()
        document = Document()
        MeetingDocConverter._configure_document_styles(document)
        lines = str(md_content or "").splitlines()
        index = 0
        while index < len(lines):
            raw = lines[index].rstrip()
            line = raw.strip()
            if not line:
                index += 1
                continue

            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading:
                level = min(len(heading.group(1)), 4)
                paragraph = document.add_heading(heading.group(2).strip(), level=level)
                MeetingDocConverter._apply_paragraph_font(
                    paragraph,
                    size_pt=MeetingDocConverter.HEADING_SIZES_PT.get(level, 12),
                    bold=True,
                )
                index += 1
                continue

            if line.startswith("|") and "|" in line[1:]:
                table_rows = [MeetingDocConverter._split_table_row(line)]
                index += 1
                if index < len(lines) and MeetingDocConverter._is_table_separator(lines[index].strip()):
                    index += 1
                while index < len(lines):
                    next_line = lines[index].strip()
                    if not next_line.startswith("|") or "|" not in next_line[1:]:
                        break
                    table_rows.append(MeetingDocConverter._split_table_row(next_line))
                    index += 1
                MeetingDocConverter._add_table(document, table_rows)
                continue

            bullet = re.match(r"^[-*]\s+(.+)$", line)
            if bullet:
                MeetingDocConverter._add_paragraph(document, bullet.group(1), style="List Bullet")
                index += 1
                continue

            numbered = re.match(r"^\d+[.、]\s+(.+)$", line)
            if numbered:
                MeetingDocConverter._add_paragraph(document, numbered.group(1), style="List Number")
                index += 1
                continue

            MeetingDocConverter._add_paragraph(document, line)
            index += 1

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def md_file_to_docx(md_path: str, docx_path: str):
        with open(md_path, "r", encoding="utf-8") as file:
            content = file.read()
        data = MeetingDocConverter.md_to_docx_bytes(content)
        os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
        with open(docx_path, "wb") as file:
            file.write(data)

    @staticmethod
    def docx_to_md_text(docx_path: str, promote_plain_headings: bool = False) -> str:
        Document = MeetingDocConverter._document_class()
        document = Document(docx_path)
        lines = []
        used_plain_headings = set()
        current_plain_section = ""
        paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else ""
            heading_level = MeetingDocConverter._docx_heading_level(paragraph)
            if heading_level is not None:
                level = max(1, min(heading_level, 6))
                lines.append(f"{chr(35) * level} {text}")
                if level == 1:
                    current_plain_section = ""
                elif level == 2:
                    current_plain_section = MeetingDocConverter._normalize_plain_heading(text)
                continue

            promoted = MeetingDocConverter._plain_heading_markdown(text) if promote_plain_headings else None
            if promoted and promoted not in used_plain_headings:
                lines.append(promoted)
                used_plain_headings.add(promoted)
                if promoted.startswith("## "):
                    current_plain_section = MeetingDocConverter._normalize_plain_heading(text)
                elif promoted.startswith("# "):
                    current_plain_section = ""
                continue

            next_paragraph = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
            if (
                promote_plain_headings
                and current_plain_section == "讨论事项综述"
                and MeetingDocConverter._is_plain_numbered_subheading(paragraph, next_paragraph)
            ):
                lines.append(f"### {text}")
            elif "List Bullet" in style:
                lines.append(f"- {text}")
            elif "List Number" in style:
                lines.append(f"1. {text}")
            else:
                lines.append(text)
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines).strip() + ("\n" if lines else "")
