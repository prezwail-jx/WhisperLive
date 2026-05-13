#!/usr/bin/env python3
import argparse
import json
from pathlib import Path


def format_range(segment):
    start = segment.get("start", "")
    end = segment.get("end", "")
    return f"{start} - {end}".strip(" -")


def segment_key(segment):
    return segment.get("start", ""), segment.get("end", "")


def load_log(path):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def add_metadata(document, data):
    document.add_heading("会议同传记录", level=0)
    metadata = [
        ("会议 ID", data.get("meeting_id", "")),
        ("创建时间", data.get("created_at", "")),
        ("导出时间", data.get("exported_at", "")),
        ("Server", data.get("server", "")),
        ("模型", data.get("model", "")),
        ("翻译模式", data.get("translation_mode", "")),
    ]
    table = document.add_table(rows=0, cols=2)
    for label, value in metadata:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value or "")


def add_segments(document, source_segments, translation_segments):
    translations = {segment_key(seg): seg for seg in translation_segments}
    source_keys = [segment_key(seg) for seg in source_segments]
    translation_only_keys = [
        key for key in translations.keys()
        if key not in set(source_keys)
    ]

    document.add_heading("正文", level=1)
    for source in source_segments:
        key = segment_key(source)
        translation = translations.get(key, {})
        document.add_paragraph(format_range(source), style=None)
        document.add_paragraph(f"原文：{source.get('text', '')}")
        document.add_paragraph(f"翻译：{translation.get('text', '')}")

    for key in translation_only_keys:
        translation = translations[key]
        document.add_paragraph(format_range(translation), style=None)
        document.add_paragraph("原文：")
        document.add_paragraph(f"翻译：{translation.get('text', '')}")


def main():
    parser = argparse.ArgumentParser(description="Convert WhisperLive meeting log JSON to Word docx.")
    parser.add_argument("input", help="Meeting log JSON exported from the Web frontend.")
    parser.add_argument("--output", "-o", default="meeting.docx", help="Output docx path.")
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError as error:
        raise SystemExit(
            "Missing dependency: python-docx. Run: pip install -r requirements/docx.txt"
        ) from error

    data = load_log(args.input)
    document = Document()
    add_metadata(document, data)
    add_segments(
        document,
        data.get("source_segments", []),
        data.get("translation_segments", []),
    )
    document.save(args.output)
    print(f"Saved: {args.output}")


if __name__ == "__main__":
    main()
